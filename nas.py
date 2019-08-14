import subprocess
import sys

def install(package):
    subprocess.call([sys.executable, "-m", "pip", "install", package])

if __name__ == '__main__':
    install('lxml')
    install('python-docx')
    install('requests')

import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from lxml import etree, html
from datetime import datetime
import requests
import re
import unicodedata
import sys
import html as pythonhtml

SINGLE_LINE_RE = re.compile('[\\n\\t\\r]|[ ]{2,}|&nbsp;')
EXTRACT_DT_RE = re.compile('^.*?(\d{2}[ ][A-Z][a-z]{2}[ ]\d{4}).*$')
ARTICLE_TYPES_MAP = {
    'NEWS RELEASES': '001',
    'SPEECHES': '002',
    'OTHERS': '003',
}
TITLE_STYLE = 'TITLE_STYLE'
DATETIME_STYLE = 'DATETIME_STYLE'
CAPTION_STYLE = 'CAPTION_STYLE'
BODY_STYLE = 'BODY_STYLE'
RUN_BODY_STYLE = 'RUN_BODY_STYLE'
MORE_RESOURCES_TITLE = 'More Resources'
MORE_RESOURCES_TITLE_STYLE = 'MORE_RESOURCES_TITLE_STYLE'
MORE_RESOURCES_LINK_STYLE = 'MORE_RESOURCES_LINK_STYLE'
LIST_BULLET_STYLE = 'LIST_BULLET_STYLE'
RUN_LIST_BULLET_STYLE = 'RUN_LIST_BULLET_STYLE'
LIST_NUMBER_STYLE = 'LIST_NUMBER_STYLE'
RUN_LIST_NUMBER_STYLE = 'RUN_LIST_NUMBER_STYLE'
TABLE_STYLE = 'TABLE_STYLE'
RUN_TABLE_STYLE = 'RUN_TABLE_STYLE'

FONT_TNR = 'Times New Roman'
LOGO_URL = 'https://www.mindef.gov.sg/web/mindefstatic/themes/Portal8.5/images/logo_mindef.png'
VISITED_MAP = dict()
DEFAULT_IMAGE_WIDTH = Cm(15.24)
LOGO_FILENAME = ''

def docx_add_bold(style):
    style.font.bold = True
    style.font.cs_bold = True

def docx_add_italic(style):
    style.font.italic = True
    style.font.cs_italic = True

def docx_add_underline(style):
    style.font.underline = True

def docx_init_styles(styles):
    title_style = styles.add_style(TITLE_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = FONT_TNR
    title_style.font.size = Pt(14)
    docx_add_bold(title_style)
    
    datetime_style = styles.add_style(DATETIME_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    datetime_style.font.name = FONT_TNR
    datetime_style.font.size = Pt(10)
    
    caption_style = styles.add_style(CAPTION_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    caption_style.font.name = FONT_TNR
    caption_style.font.size = Pt(10)
    
    body_style = styles.add_style(BODY_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = FONT_TNR
    body_style.font.size = Pt(12)
    
    run_body_style = styles.add_style(RUN_BODY_STYLE, WD_STYLE_TYPE.CHARACTER)
    run_body_style.base_style = styles[BODY_STYLE]
    
    list_bullet_style = styles.add_style(LIST_BULLET_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    list_bullet_style.base_style = styles['List Bullet']
    list_bullet_style.font.name = FONT_TNR
    list_bullet_style.font.size = Pt(12)
    
    run_list_bullet_style = styles.add_style(RUN_LIST_BULLET_STYLE, WD_STYLE_TYPE.CHARACTER)
    run_list_bullet_style.base_style = styles[LIST_BULLET_STYLE]
    
    list_number_style = styles.add_style(LIST_NUMBER_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    list_number_style.base_style = styles['List Number']
    list_number_style.font.name = FONT_TNR
    list_number_style.font.size = Pt(12)
    
    run_list_number_style = styles.add_style(RUN_LIST_NUMBER_STYLE, WD_STYLE_TYPE.CHARACTER)
    run_list_number_style.base_style = styles[LIST_NUMBER_STYLE]
    print(dir(styles[LIST_NUMBER_STYLE]), styles[LIST_NUMBER_STYLE].element)
    
    more_resources_title_style = styles.add_style(MORE_RESOURCES_TITLE_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    more_resources_title_style.base_style = styles[BODY_STYLE]
    docx_add_bold(more_resources_title_style)
    
    more_resources_link_style = styles.add_style(MORE_RESOURCES_LINK_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    more_resources_link_style.font.name = FONT_TNR
    more_resources_link_style.font.size = Pt(10)
    
    table_style = styles.add_style(TABLE_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    table_style.font.name = FONT_TNR
    table_style.font.size = Pt(12)
    table_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run_table_style = styles.add_style(RUN_TABLE_STYLE, WD_STYLE_TYPE.CHARACTER)
    run_table_style.base_style = styles[TABLE_STYLE]


def parse_cleanup(txt, dont_trim=False):
#     return unicodedata.normalize("NFKD", re.sub(SINGLE_LINE_RE, '', txt)).strip()
    clean = pythonhtml.unescape(re.sub(SINGLE_LINE_RE, '', txt))
    
    return clean.strip() if not dont_trim else clean


def parse_extract_datetime(dt):
    return re.sub(EXTRACT_DT_RE, '\\1', dt)


def parse_append_hostname(url):
    return 'https://www.mindef.gov.sg{}'.format(url) if url.startswith('/') else url


def parse_filename(dt, article_type_num='001'):
    filename = datetime.strptime(dt, '%d %b %Y')

    return '{}{}'.format(filename.strftime('%Y%m%d'), article_type_num)

        
def parse_fetch_image(url, idx):
    image_filename = 'img{}.png'.format(idx)
    image = requests.get(url)
    open(image_filename, 'wb').write(image.content)

    return image_filename


def parse_article(url):
    global VISITED_MAP

    if url in VISITED_MAP:
        return 'VISITED'
    else:
        VISITED_MAP[url] = True
        
    try:
        page = requests.get(url)
    except Exception as ex:
        print('-------------------------')
        print(url, ex)
        print('-------------------------')

        return
        
    tree = html.fromstring(page.content)
    
    titles = tree.xpath('//div[@class="article-detail__heading"]/div[contains(@class, "title")]/text()')
    title = parse_cleanup(titles[0])
    
    article_types = tree.xpath('//div[@class="article-detail__heading"]/div[@class="article-info"]/span[contains(@class, "item-label")]/text()')
    article_type = parse_cleanup(article_types[0]).upper()
    
    datetime_strs = tree.xpath('//div[@class="article-detail__heading"]/div[@class="article-info"]/span[contains(@class, "item-published")]/text()')
    datetime_str = parse_extract_datetime(datetime_strs[0])
    
    images = tree.xpath('//article[contains(@class, "mindef-gallery-container")]/div[contains(@class, "mindef-gallery")]//img/@src')
    images = list(map(parse_append_hostname, images))
    for idx, img_url in enumerate(images):
        img_filename = parse_fetch_image(img_url, idx)
        images[idx] = img_filename
        
    captions = tree.xpath('//article[contains(@class, "mindef-gallery-container")]/div[contains(@class, "mindef-gallery")]//span[@class="caption"]/text()')
    captions = list(map(parse_cleanup, captions))
    
    body = tree.xpath('//div[@class="row"]//div[@class="article"]')[0]
    # body = list(filter(lambda v: len(v) > 0, list(map(parse_cleanup, body))))
#     for idx, b in enumerate(body):
#         body[idx] = etree.tostring(b, pretty_print=True)

    others_text = tree.xpath('//div[@class="more-resources"]/div[@class="more-resources__links"]/p/a/text()')
    others_text = list(map(parse_cleanup, others_text))
    
    others_link = tree.xpath('//div[@class="more-resources"]/div[@class="more-resources__links"]/p/a/@href')
    others_link = list(map(parse_append_hostname, others_link))

    print('Title', title)
    print('Article Type', article_type, ARTICLE_TYPES_MAP[article_type])
    print('DateTime', datetime_str)
    
    print('Images', images)
    
    print('Captions', captions)
    
    print('Body', body)
    
    print('Others Text', others_text)
    print('Others Link', others_link)
    
    docx_build(article_type, title, datetime_str, images, captions, body, others_text, others_link)
    
    for ol in others_link:
        parse_article(ol)

"""
Source: https://github.com/python-openxml/python-docx/issues/384
:param paragraph: The paragraph we are adding the hyperlink to.
:param url: A string containing the required url
:param text: The text displayed for the url
:return: The hyperlink object
"""
def docx_add_hyperlink(paragraph, url, text):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


"""
Source: https://stackoverflow.com/questions/51829366/bullet-lists-in-python-docx
Makes a paragraph into a list item with a specific level and
optional restart.

An attempt will be made to retreive an abstract numbering style that
corresponds to the style of the paragraph. If that is not possible,
the default numbering or bullet style will be used based on the
``num`` parameter.

Parameters
----------
doc : docx.document.Document
    The document to add the list into.
par : docx.paragraph.Paragraph
    The paragraph to turn into a list item.
prev : docx.paragraph.Paragraph or None
    The previous paragraph in the list. If specified, the numbering
    and styles will be taken as a continuation of this paragraph.
    If omitted, a new numbering scheme will be started.
level : int or None
    The level of the paragraph within the outline. If ``prev`` is
    set, defaults to the same level as in ``prev``. Otherwise,
    defaults to zero.
num : bool
    If ``prev`` is :py:obj:`None` and the style of the paragraph
    does not correspond to an existing numbering style, this will
    determine wether or not the list will be numbered or bulleted.
    The result is not guaranteed, but is fairly safe for most Word
    templates.
"""
def docx_add_bullet(doc, par, prev=None, level=None, num=True):
    xpath_options = {
        True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
        False: {'single': '', 'level': level},
    }

    def style_xpath(prefer_single=True):
        """
        The style comes from the outer-scope variable ``par.style.name``.
        """
        style = par.style.style_id
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
            ']/@w:abstractNumId'
        ).format(style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        """
        The type is from the outer-scope variable ``num``.
        """
        type = 'decimal' if num else 'bullet'
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
            ']/@w:abstractNumId'
        ).format(type=type, **xpath_options[prefer_single])

    def get_abstract_id():
        """
        Select as follows:

            1. Match single-level by style (get min ID)
            2. Match exact style and level (get min ID)
            3. Match single-level decimal/bullet types (get min ID)
            4. Match decimal/bullet in requested level (get min ID)
            3. 0
        """
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if (prev is None or
            prev._p.pPr is None or
            prev._p.pPr.numPr is None or
            prev._p.pPr.numPr.numId is None):
        if level is None:
            level = 0
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        # Compute the abstract ID first by style, then by num
        anum = get_abstract_id()
        # Set the concrete numbering based on the abstract numbering ID
        num = numbering.add_num(anum)
        # Make sure to override the abstract continuation property
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        # Extract the newly-allocated concrete numbering ID
        num = num.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        # Get the previous concrete numbering ID
        num = prev._p.pPr.numPr.numId.val
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level


def docx_get_colwidths(tablerow):
    colwidths = []

    for col in tablerow.getchildren():
        if col.tag == 'td':
            if 'colspan' not in col.attrib:
                colwidth = int(col.attrib['width']) if 'width' in col.attrib and col.attrib['width'] != '' else None
                colwidths.append(colwidth)
            else:
                colspan = int(col.attrib['colspan']) if col.attrib['colspan'] != '' else 1
                for i in range(colspan):
                    colwidths.append(None)

    return colwidths


def docx_build_rows_cols(tablebody, docx):
    numtablerows = 0
    colwidths = []

    for tablerow in tablebody.getchildren():
        tag = tablerow.tag
        
        if tag == 'tbody':
            for tablerow_ in tablerow.getchildren():
                if tablerow_.tag == 'tr':
                    numtablerows += 1
                    
                    if len(colwidths) > 0:
                        continue
                    
                    colwidths = docx_get_colwidths(tablerow_)
        elif tag == 'tr':
            numtablerows += 1

            if len(colwidths) > 0:
                continue

            colwidths = docx_get_colwidths(tablerow)
    
    numtablecols = len(colwidths)
    print('rows', numtablerows, 'cols', numtablecols)
    docxtable = docx.add_table(rows=numtablerows, cols=numtablecols)
    default_width = Pt(numtablecols * 0.01)

    for colidx, colwidth in enumerate(colwidths):
        if colwidth is not None:
            docxtable.columns[colidx].width = Pt(colwidth)
        else:
            docxtable.columns[colidx].width = default_width

    return docxtable


def docx_build_row_data(rowidx, tablerow, docxtable):
    colidx = 0
    tablerow_children = tablerow.getchildren()

    for tablecol in tablerow_children:
        print('loop', colidx)
        coltag = tablecol.tag
        
        if coltag == 'td':
            if 'colspan' in tablecol.attrib:
                colspan = int(tablecol.attrib['colspan'])
                print('colspan', colspan, colidx)
                docxcell = docxtable.cell(rowidx, colidx)
                docxcell.style = TABLE_STYLE
                endcolidx = colidx + colspan
                colidx += 1
                while colidx < endcolidx:
                    print('colspan', colidx)
                    docxcell = docxcell.merge(docxtable.cell(rowidx, colidx))
                    colidx += 1
                print('colspan', colidx)
            else:
                print('else', colidx)
                docxcell = docxtable.cell(rowidx, colidx)
                docxcell.style = TABLE_STYLE
                colidx += 1
                print('else', colidx)
            print('Build Cell')
            tablecol_text = tablecol.text

            if tablecol_text is not None and tablecol_text != '':
                tablecol_text = parse_cleanup(tablecol.text, dont_trim=True)
                docxcell.add_paragraph(tablecol_text, style=TABLE_STYLE)
            if len(tablecol.getchildren()) > 0:
                docx_build_body(tablecol, doc=docxcell)



def docx_build_table(tablebody, docx):
    print('Build Table')
    docxtable = docx_build_rows_cols(tablebody, docx)
    docxtable.alignment = WD_TABLE_ALIGNMENT.CENTER

    for rowidx, tablerow in enumerate(tablebody.getchildren()):
        rowtag = tablerow.tag
        
        if rowtag == 'tbody':
            for rowidx_, tablerow_ in enumerate(tablerow.getchildren()):
                if tablerow_.tag == 'tr':
                    docx_build_row_data(rowidx_, tablerow_, docxtable)
        elif rowtag == 'tr':
            docx_build_row_data(rowidx, tablerow, docxtable)
            

def docx_build_body(body, doc=None, paragraph=None, run=None, parent_tag='', parent_run_style=RUN_BODY_STYLE):
    para_style = TABLE_STYLE if body.tag == 'td' else BODY_STYLE

    for e in body.getchildren():
        tag = e.tag

        if tag == 'table':
            docx_build_table(e, doc)
            continue

        para_to_use = doc.add_paragraph('', style=para_style) if paragraph is None else paragraph
        before_text = parse_cleanup(e.text if e.text is not None else '', dont_trim=True)
        after_text = parse_cleanup(e.tail if e.tail is not None else '', dont_trim=True)
        before_run = run
        run_style = parent_run_style
        print('INFO', parent_tag, tag, '|'+before_text, '|'+after_text, before_run, run_style)

        # Start of any element
        if tag == 'a' and 'href' in e.attrib and e.attrib['href'] is not None and e.attrib['href'] != '':
            url = e.attrib['href']
            docx_add_hyperlink(para_to_use, url, before_text)
        elif tag == 'li':
            print('LI')
            before_run = para_to_use.add_run('', style=run_style)
            before_run.add_break(WD_BREAK.LINE)
            before_run.add_break(WD_BREAK.LINE)
        elif before_run is None:
            print('NEW RUN')
            if tag == 'ul':
                run_style = RUN_LIST_BULLET_STYLE
            if tag == 'ol':
                run_style = RUN_LIST_NUMBER_STYLE
            before_run = para_to_use.add_run('', style=run_style)
        else:
            print('OLD RUN')

        if tag == 'br':
            before_run.add_break(WD_BREAK.LINE)
        elif tag == 'hr':
            before_run.add_break(WD_BREAK.PAGE)
        elif tag != 'a' and before_text != '':
            before_run.text = before_text

        if tag == 'b' or tag == 'strong':
            docx_add_bold(before_run)
        elif tag == 'i' or tag == 'italic' or tag == 'em':
            docx_add_italic(before_run)
        elif tag == 'u':
            docx_add_underline(before_run)
        elif tag == 'h1' or tag == 'h2' or tag == 'h3' or tag == 'h4' or tag == 'h5'or tag == 'h6':
            before_run.font.size = Pt(12)
            docx_add_bold(before_run)

        if after_text is not '':
            if run is None:
                # print('AFTERTEXT NONE RUN BEFORE TEXT', after_text)
                para_to_use.add_run(after_text, style=RUN_BODY_STYLE)
            else:
                rbt = before_run.text
                before_run.text = '{}{}'.format(rbt if rbt is not None else '', after_text)
                # print('AFTERTEXT RUN BEFORE TEXT', before_run.text)

        # Only pass run if its not a paragraph tag
        docx_build_body(e, doc=doc, paragraph=para_to_use, run=before_run if tag != 'p' else None, parent_tag=tag, parent_run_style=run_style)


    return []


"""
https://www.mindef.gov.sg/web/wcm/connect/mindef/mindef-content/home?siteAreaName=&srv=cmpnt&selectedCategories=news-releases&cmpntid=dcb39e68-0637-4383-b587-29be9bb9bea5&source=library&cache=none&contentcache=none&connectorcache=none&wcm_page.MENU-latest-releases=3
"""
def docx_build(article_type, title, datetime_str, images, captions, body, others_text, others_link):
    dup_prefix = ''
    filename = parse_filename(datetime_str, ARTICLE_TYPES_MAP[article_type])
    print(filename)
    filename = 'MINDEF_{}{}.docx'.format(filename , dup_prefix)
    print(filename)
    doc = Document()
    docx_init_styles(doc.styles)
    
    doc.add_picture(LOGO_FILENAME, width=DEFAULT_IMAGE_WIDTH)
    
    doc.add_paragraph(title, style=TITLE_STYLE)
    
    doc.add_paragraph(datetime_str, style=DATETIME_STYLE)
    
    num_images = len(images)
    num_captions = len(captions)
    num_overall = num_images if num_images > num_captions else num_captions
    
    if num_images > 0:
        doc.add_picture(images[0], width=DEFAULT_IMAGE_WIDTH)
    
    if num_captions > 0:
        doc.add_paragraph(captions[0], style=CAPTION_STYLE)
    
    docx_build_body(body, doc)

    for i in range(1, num_overall):
        doc.add_picture(images[i], width=DEFAULT_IMAGE_WIDTH)
        doc.add_paragraph(captions[i], style=CAPTION_STYLE)

    num_texts = len(others_text)
    num_links = len(others_link)
    num_overall = num_texts if num_texts > num_links else num_links
    
    if num_overall > 0:
        doc.add_paragraph(MORE_RESOURCES_TITLE, style=MORE_RESOURCES_TITLE_STYLE)
        for i in range(num_overall):
            other_para = doc.add_paragraph('', style=MORE_RESOURCES_LINK_STYLE)
            docx_add_hyperlink(other_para, others_link[i], others_text[i])

    doc.save(filename)

def docx_test():
    xxx = html.fromstring('<html><head><title>sadasdas</title></head><body>\
        <div style="overflow:auto;"><table border="1" dir="ltr" style="width: 100%; border-collapse : collapse; border-color : #696969;">\
    <tbody>\
        <tr>\
            <td colspan="3" style="border-color: rgb(105, 105, 105); text-align: center;"><strong>Jury\'s Choice Awards</strong></td>\
        </tr>\
        <tr>\
            <td colspan="3" style="border-color: rgb(105, 105, 105); text-align: center;"><em><strong>The winners for the following awards were determined by a selection panel set up by the Organiser.</strong></em></td>\
        </tr>\
        <tr>\
            <td style="border-color: rgb(105, 105, 105); text-align: center; width: 223px;"><strong>Prizes</strong></td>\
            <td style="border-color: rgb(105, 105, 105); text-align: center; width: 328px;"><strong>Student Category</strong></td>\
            <td style="border-color: rgb(105, 105, 105); text-align: center;"><strong>Open Category</strong></td>\
        </tr>\
        <tr>\
            <td style="border-color: rgb(105, 105, 105); width: 223px;">\
            <ul>\
                <li><strong>Feature Film Deal</strong></li>\
            </ul>\
            </td>\
            <td colspan="2" rowspan="1" style="border-color: rgb(105, 105, 105); width: 579px; text-align: center;">ciNE65 Movie Makers Award<br>\
            <br>\
            <em><strong>Winner:</strong> My Homeland: A Photography Project by Grandpa Chen</em></td>\
        </tr>\
        <tr>\
            <td style="border-color: rgb(105, 105, 105); width: 223px;">\
            <ul>\
                <li><strong>Cash Prize $3,000</strong><br>\
                &nbsp;</li>\
                <li><strong>Learning Trip to an International Film Festival</strong><br>\
                &nbsp;</li>\
                <li><strong>Panasonic Professional 4K Camcorder (AG-UX180EN)&nbsp;</strong></li>\
            </ul>\
            </td>\
            <td style="border-color: rgb(105, 105, 105); width: 328px;">\
            <p style="text-align: center;">Overall Best Film Award</p>\
\
            <p style="text-align: center;"><em><strong>Winner: </strong>My Homeland: A Photography Project<br>\
            by Grandpa Chen</em></p>\
            </td>\
            <td style=" border-color : #696969;">\
            <p style="text-align: center;">Overall Best Film Award</p>\
\
            <p style="text-align: center;"><em><strong>Winner: </strong>$ingapura</em></p>\
            </td>\
        </tr>\
        <tr>\
            <td style="border-color: rgb(105, 105, 105); width: 223px;">\
            <ul>\
                <li><strong>Cash Prize $1,000&nbsp;</strong><br>\
                &nbsp;</li>\
                <li><strong>Panasonic Professional 4K Camera (AG-UX90EN)</strong></li>\
            </ul>\
            </td>\
            <td style="border-color: rgb(105, 105, 105); width: 328px;">\
            <p style="text-align: center;">Best Cinematography</p>\
\
            <p style="text-align: center;"><em><strong>Winner:</strong> 村 Kampong</em></p>\
            </td>\
            <td style=" border-color : #696969;">\
            <p style="text-align: center;">Best Cinematography</p>\
\
            <p style="text-align: center;"><em><strong>Winner: </strong>一人一半</em></p>\
            </td>\
        </tr>\
        <tr>\
            <td colspan="1" rowspan="5" style="border-color: rgb(105, 105, 105); width: 223px;">\
            <ul>\
                <li><strong>Cash Prize $1,000&nbsp;</strong><br>\
                &nbsp;</li>\
                <li><strong>Panasonic product voucher worth $288 per award</strong></li>\
            </ul>\
            </td>\
            <td style="border-color: rgb(105, 105, 105); width: 328px; text-align: center;">Best Direction<br>\
            <br>\
            <em><strong>Winner: </strong>Broken</em></td>\
            <td style="border-color: rgb(105, 105, 105); text-align: center;">Best Direction<br>\
            <br>\
            <em><strong>Winner: </strong>一人一半</em></td>\
        </tr>\
        <tr>\
            <td style="border-color: rgb(105, 105, 105); width: 328px; text-align: center;">Best Editing<br>\
            <br>\
            <em><strong>Winner: </strong>Echoes of 1965</em></td>\
            <td style="border-color: rgb(105, 105, 105); text-align: center;">Best Editing<br>\
            <br>\
            <em><strong>Winner: </strong>Ah Gong Garden</em></td>\
        </tr>\
        <tr>\
            <td style="border-color: rgb(105, 105, 105); width: 328px; text-align: center;">Best Screenplay<br>\
            <br>\
            <em><strong>Winner:</strong> Built to Last</em></td>\
            <td style="border-color: rgb(105, 105, 105); text-align: center;">Best Screenplay<br>\
            <br>\
            <em><strong>Winner: </strong>$ingapura</em></td>\
        </tr>\
        <tr>\
            <td style="border-color: rgb(105, 105, 105); width: 328px; text-align: center;">Best Sound Design<br>\
            <br>\
            <em><strong>Winner: </strong>Pulau Ujong</em></td>\
            <td style="border-color: rgb(105, 105, 105); text-align: center;">Best Sound Design<br>\
            <br>\
            <em><strong>Winner: </strong>Sound of Home</em></td>\
        </tr>\
        <tr>\
            <td style="border-color: rgb(105, 105, 105); width: 328px; text-align: center;">Best Art Direction<br>\
            <br>\
            <em><strong>Winner: </strong>chope</em></td>\
            <td style="border-color: rgb(105, 105, 105); text-align: center;">Best Art Direction<br>\
            <br>\
            <em><strong>Winner: </strong>一人一半</em></td>\
        </tr>\
    </tbody>\
</table></div>\
        <p><b><i><u>hypernest</u></i></b></p>\
        <p>ptext before<b>ptext bold</b><i>ptext italic</i>ptext after</p>\
        <p>ptext bold before<b>bold</b>ptext bold after</p>\
        <p><b>asdad</b> xxx <strong>strong</strong> xxx <b>sadas</b><br />xxx<br>xxx<i>italic</i>xxx<u>underline</u>xxx<ol><li>num 1</li><li>num 2</li></ol><ul><li>bullet 1</li><li>bullet 2</li></ul></p></body></html>')
    body = xxx.xpath('//body')[0]
    doc = Document()
    docx_init_styles(doc.styles)
    docx_build_body(body, doc)
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc
    doc.save('test.docx')
    # para = doc.add_paragraph('', style = BODY_STYLE)
    # txt = para.add_run('asdasdas', style = RUN_BODY_STYLE)
    # docx_add_bold(txt)
    # docx_add_italic(txt)
    # docx_add_underline(txt)
    # txt = para.add_run('XXXXXX', style = RUN_BODY_STYLE)
    # para = doc.add_paragraph('A', style=LIST_NUMBER_STYLE)
    # para = doc.add_paragraph('B', style=LIST_NUMBER_STYLE)
    # para = doc.add_paragraph('C', style=LIST_NUMBER_STYLE)
    # para = doc.add_paragraph('A', style=LIST_BULLET_STYLE)
    # para = doc.add_paragraph('B', style=LIST_BULLET_STYLE)
    # para = doc.add_paragraph('C', style=LIST_BULLET_STYLE)
    # other_para = doc.add_paragraph('', style=LIST_BULLET_STYLE)
    # docx_add_bold(other_para.style)
    # docx_add_italic(other_para.style)
    # docx_add_underline(other_para.style)
    # docx_add_hyperlink(other_para, 'http://www.google.com.sg', 'Testing')
    # other_para = doc.add_paragraph('', style=BODY_STYLE)
    # test_run = other_para.add_run('\r\nasdas', style=RUN_LIST_BULLET_STYLE)
    # test_run = other_para.add_run('\r\nasdas', style=RUN_LIST_NUMBER_STYLE)
    # doc.save('test.docx')

# docx_test()

def get_pages(category, year, long_month, page):
    url = 'https://www.mindef.gov.sg/web/wcm/connect/mindef/mindef-content/home'
    params = {
        'siteAreaName': 'mindef-content/home/news-and-events/latest-releases/{}/{}'.format(year, long_month),
        'srv': 'cmpnt',
        'selectedCategories': category,
        'cmpntid': 'dcb39e68-0637-4383-b587-29be9bb9bea5',
        'source': 'library',
        'cache': 'none',
        'contentcache': 'none',
        'connectorcache': 'none',
        'wcm_page.MENU-latest-releases': page,
    }
    page = requests.get(url, params)
    if page.content == None or page.content == '':
        return None

    tree = html.fromstring(page.content)
    
    links = tree.xpath('//a[@class="news-event-item-link"]/@href')
    links = list(map(parse_append_hostname, links))

    datetimes = tree.xpath('//span[@class="type-body-2"]/text()')
    print(datetimes[0])
    datetimes = list(map(parse_extract_datetime, datetimes))
    print(datetimes[0])
    datetimes = list(map(parse_filename, datetimes))

    return {
        'links': links,
        'datetimes': datetimes,
    }

def get_month_pages(category, year, long_month):
    month = {
        'links': [],
        'datetimes': [],
    }

    for page in range(1, 100):
        pages = get_pages(category, year, long_month, page)
        if pages == None:
            break
        month.links = month.links + pages.links
        month.datetimes = month.datetimes + pages.datetimes

    return month

def get_year_pages(category, year):
    year = {
        'links': [],
        'datetimes': [],
    }

    for month_str in ['january', 'february', 'march', 'april', 'may', 'june', 'july', 'august', 'september', 'october', 'november', 'december']:
        month = get_month_pages(category, year, month_str)
        year.links = year.links + month.links
        year.datetimes = year.datetimes + month.datetimes

    return year



# LOGO_FILENAME = parse_fetch_image(LOGO_URL, 'LOGO')
# parse_article(url='https://www.mindef.gov.sg/web/portal/mindef/news-and-events/latest-releases/article-detail/2019/june/13jun19_nr')
print(get_year_pages('new-releases', 2013))