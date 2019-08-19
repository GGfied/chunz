import docx
from builddocx_body_table import docx_build_body
from constants import *
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt
from docx_helpers import docx_add_bold, docx_add_hyperlink, docx_add_underline, docx_apply_hyperlink_style, \
    docx_add_italic
from globals import GLOBALS
from parse_helpers import parse_fetch_image
from writers import write_error


def docx_init_styles(styles):
    title_style = styles.add_style(TITLE_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = FONT_TNR
    title_style.font.size = Pt(14)
    docx_add_bold(title_style)

    datetime_style = styles.add_style(DATETIME_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    datetime_style.font.name = FONT_TNR
    datetime_style.font.size = Pt(10)

    caption_style = styles.add_style(CAPTION_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    caption_style.font.name = FONT_TNR
    caption_style.font.size = Pt(10)
    docx_add_italic(caption_style)

    run_caption_style = styles.add_style(RUN_CAPTION_STYLE, WD_STYLE_TYPE.CHARACTER)
    run_caption_style.font.name = FONT_TNR
    run_caption_style.font.size = Pt(10)
    docx_add_italic(run_caption_style)

    body_style = styles.add_style(BODY_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = FONT_TNR
    body_style.font.size = Pt(12)

    run_body_style = styles.add_style(RUN_BODY_STYLE, WD_STYLE_TYPE.CHARACTER)
    run_body_style.base_style = styles[BODY_STYLE]

    list_bullet_style = styles.add_style(LIST_BULLET_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    list_bullet_style.base_style = styles['List Bullet']
    list_bullet_style.font.name = FONT_TNR
    list_bullet_style.font.size = Pt(12)

    run_list_bullet_style = styles.add_style(RUN_LIST_BULLET_STYLE, WD_STYLE_TYPE.CHARACTER)
    run_list_bullet_style.base_style = styles[LIST_BULLET_STYLE]

    list_number_style = styles.add_style(LIST_NUMBER_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    list_number_style.base_style = styles['List Number']
    list_number_style.font.name = FONT_TNR
    list_number_style.font.size = Pt(12)

    run_list_number_style = styles.add_style(RUN_LIST_NUMBER_STYLE, WD_STYLE_TYPE.CHARACTER)
    run_list_number_style.base_style = styles[LIST_NUMBER_STYLE]

    more_resources_title_style = styles.add_style(MORE_RESOURCES_TITLE_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    more_resources_title_style.base_style = styles[BODY_STYLE]
    docx_add_bold(more_resources_title_style)

    more_resources_link_style = styles.add_style(MORE_RESOURCES_LINK_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    more_resources_link_style.font.name = FONT_TNR
    more_resources_link_style.font.size = Pt(10)
    docx_add_underline(more_resources_link_style)

    table_style = styles.add_style(TABLE_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    table_style.base_style = styles[BODY_STYLE]

    run_table_style = styles.add_style(RUN_TABLE_STYLE, WD_STYLE_TYPE.CHARACTER)
    run_table_style.base_style = styles[TABLE_STYLE]

    run_link_style = styles.add_style(RUN_LINK_STYLE, WD_STYLE_TYPE.CHARACTER)
    run_link_style.base_style = styles[RUN_BODY_STYLE]
    docx_apply_hyperlink_style(run_link_style)


"""
https://www.mindef.gov.sg/web/wcm/connect/mindef/mindef-content/home?siteAreaName=&srv=cmpnt&selectedCategories=news-releases&cmpntid=dcb39e68-0637-4383-b587-29be9bb9bea5&source=library&cache=none&contentcache=none&connectorcache=none&wcm_page.MENU-latest-releases=3
"""


def docx_build(save_filename, filename_prefix, directory, title, datetime_str, images, body, others_text, others_link):
    doc = Document()
    docx_init_styles(doc.styles)

    doc.add_picture(GLOBALS['LOGO_FILENAME'], width=DEFAULT_IMAGE_WIDTH)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph(title, style=TITLE_STYLE)

    doc.add_paragraph(datetime_str, style=DATETIME_STYLE)

    for idx, img in enumerate(images):
        img_filename = parse_fetch_image(img['link'], idx, filename_prefix, directory)
        images[idx]['link'] = img_filename
    num_images = len(images)

    if num_images > 0:
        try:
            doc.add_picture(images[0]['link'], width=DEFAULT_IMAGE_WIDTH)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.paragraphs[-1].runs[0].add_break(WD_BREAK.LINE)
            doc.paragraphs[-1].add_run(DEFAULT_CAPTION_PADDING + images[0]['caption'], style=RUN_CAPTION_STYLE)
        except docx.image.exceptions.UnrecognizedImageError as ex:
            write_error(directory, error='Image Error', exception=ex)
            doc.add_picture(GLOBALS['LOGO_FILENAME'], width=DEFAULT_IMAGE_WIDTH)

    docx_build_body(body, doc)

    for i in range(1, num_images):
        try:
            doc.add_picture(images[i]['link'], width=DEFAULT_IMAGE_WIDTH)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.paragraphs[-1].runs[0].add_break(WD_BREAK.LINE)
            doc.paragraphs[-1].add_run(DEFAULT_CAPTION_PADDING + images[i]['caption'], style=RUN_CAPTION_STYLE)
        except docx.image.exceptions.UnrecognizedImageError as ex:
            write_error(directory, error='Image Error', exception=ex)
            doc.add_picture(GLOBALS['LOGO_FILENAME'], width=DEFAULT_IMAGE_WIDTH)

    num_texts = len(others_text)
    num_links = len(others_link)
    num_overall = num_texts if num_texts > num_links else num_links

    if num_overall > 0:
        doc.add_paragraph(MORE_RESOURCES_TITLE, style=MORE_RESOURCES_TITLE_STYLE)
        for i in range(num_overall):
            other_para = doc.add_paragraph('', style=MORE_RESOURCES_LINK_STYLE)
            docx_add_hyperlink(other_para, others_link[i], others_text[i])

    save_path = os.path.join(directory, save_filename)
    doc.save(save_path)
    # docxtopdf.convert_to(folder=directory, source=save_path)

    return save_filename
