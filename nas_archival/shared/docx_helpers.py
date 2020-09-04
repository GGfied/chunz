import random
import re
import string
import traceback
from copy import deepcopy

import docx
from PIL import Image, ImageFile
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, RGBColor
from shared.constants import WIDTH_ATTRIB, COLSPAN_ATTRIB, RUN_LINK_STYLE, STYLE_TAG, ROWSPAN_ATTRIB, HEIGHT_ATTRIB, \
    TEXT_ALIGN_IN_STYLE_RE, \
    HEIGHT_IN_STYLE_RE, WIDTH_IN_STYLE_RE, BODY_STYLE, GLOBAL_LOGO_FILENAME, DEFAULT_IMAGE_WIDTH, \
    DEFAULT_CAPTION_PADDING, RUN_CAPTION_STYLE, TABLECELL_DIMENS_HEIGHT, TABLECELL_DIMENS_WIDTH, \
    TABLECELL_DIMENS_CELL, TABLECELL_DIMENS_CELLID, EXT_DOCX, THEAD_TAG, TBODY_TAG, TR_TAG, TH_TAG, TD_TAG
from shared.globals import GLOBALS
from shared.writers import write_error


def docx_get_filename_prefix(filename, related_count=None):
    return 'MINDEF_{}_{}'.format(filename, related_count) if related_count else 'MINDEF_{}'.format(filename)


def docx_get_save_filename(filename_prefix, ext=EXT_DOCX):
    return '{}{}'.format(filename_prefix, ext)


def docx_get_others_text(text, link):
    return '{} (Document No: {})'.format(text, link)


def docx_add_bold(style):
    style.font.bold = True
    style.font.cs_bold = True


def docx_add_italic(style):
    style.font.italic = True
    style.font.cs_italic = True


def docx_add_underline(style):
    style.font.underline = True


def docx_add_superscript(style):
    style.font.superscript = True


def docx_add_subscript(style):
    style.font.subscript = True


def docx_apply_hyperlink_style(style):
    style.font.color.rgb = RGBColor(0, 0, 255)
    style.font.underline = True


def docx_add_heading(style):
    style.font.size = Pt(12)
    docx_add_bold(style)


def docx_apply_center_align(style):
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def docx_apply_text_align(para, htmlattrib):
    whichalign = 'left'

    if STYLE_TAG in htmlattrib:
        res = re.search(TEXT_ALIGN_IN_STYLE_RE, htmlattrib[STYLE_TAG])
        whichalign = res.group(1).lower() if res else whichalign

    if whichalign == 'right':
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif whichalign == 'left':
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif whichalign == 'center':
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def docx_have_hyperlink(paragraph):
    for p in paragraph._p:
        if p.tag.endswith('}hyperlink'):
            return True

    return False


"""
Source: https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx
Source 2: https://github.com/python-openxml/python-docx/issues/384
Source 3: https://github.com/python-openxml/python-docx/issues/74
:param paragraph: The paragraph we are adding the hyperlink to.
:param url: A string containing the required url
:param text: The text displayed for the url
:return: The hyperlink object
"""


def docx_add_hyperlink(paragraph, url, text, old_run=None):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    new_run = paragraph.add_run('', style=RUN_LINK_STYLE) if old_run is None else old_run
    new_run.text = new_run.text + text
    docx_apply_hyperlink_style(new_run)

    hyperlink.append(new_run._r)
    paragraph._p.append(hyperlink)

    return new_run


"""
Source: https://github.com/python-openxml/python-docx/issues/217
"""


def docx_add_listitem(paragraph, list_type):
    p = paragraph._p  # access to xml paragraph element
    pPr = p.get_or_add_pPr()  # access paragraph properties
    numPr = docx.oxml.shared.OxmlElement('w:numPr')  # create number properties element
    numId = docx.oxml.shared.OxmlElement('w:numId')  # create numId element - sets bullet type
    numId.set(docx.oxml.ns.qn('w:val'), list_type)  # set list type/indentation
    numPr.append(numId)  # add bullet type to number properties list
    pPr.append(numPr)  # add number properties to paragraph


def get_height_in_style(style):
    res = re.search(HEIGHT_IN_STYLE_RE, style)

    return res.group(1) if res else -1


def get_width_in_style(style):
    res = re.search(WIDTH_IN_STYLE_RE, style)

    return res.group(1) if res else -1


def docx_get_coldimensions(tablebody):
    tablerows = tablebody.getchildren()
    is_first_element_thead = len(tablerows) > 0 and tablerows[0].tag.lower() == THEAD_TAG
    is_first_element_tbody = len(tablerows) > 0 and tablerows[0].tag.lower() == TBODY_TAG

    if is_first_element_thead:
        temp_tablerows = tablerows[0].getchildren()
        is_second_element_tbody = len(tablerows) > 1 and tablerows[1].tag.lower() == TBODY_TAG
        if is_second_element_tbody:
            temp_tablerows.extend(tablerows[1].getchildren())
        tablerows = temp_tablerows
    elif is_first_element_tbody:
        tablerows = tablerows[0].getchildren()
    coldimens = []

    for rowidx, row in enumerate(tablerows):
        if row.tag.lower() != TR_TAG:
            continue

        if rowidx >= len(coldimens):
            coldimens.append([])
        dimencolidx = 0

        for colidx, cell in enumerate(row.getchildren()):
            dimenrowidx = rowidx
            if cell.tag.lower() != TD_TAG and cell.tag.lower() != TH_TAG:
                continue

            while dimencolidx < len(coldimens[dimenrowidx]) and coldimens[dimenrowidx][dimencolidx] is not '':
                dimencolidx += 1

            colspan = 1
            if COLSPAN_ATTRIB in cell.attrib:
                try:
                    colspan = int(cell.attrib[COLSPAN_ATTRIB])
                except ValueError:
                    colspan = 1

            width = -1
            if WIDTH_ATTRIB in cell.attrib:
                try:
                    width = int(cell.attrib[WIDTH_ATTRIB])
                except ValueError:
                    width = -1

            if width == -1 and STYLE_TAG in cell.attrib:
                width = get_width_in_style(cell.attrib[STYLE_TAG])

            rowspan = 1
            if ROWSPAN_ATTRIB in cell.attrib:
                try:
                    rowspan = int(cell.attrib[ROWSPAN_ATTRIB])
                except ValueError:
                    rowspan = 1

            height = -1
            if HEIGHT_ATTRIB in cell.attrib:
                try:
                    height = int(cell.attrib[HEIGHT_ATTRIB])
                except ValueError:
                    height = -1

            if height == -1 and STYLE_TAG in cell.attrib:
                height = get_height_in_style(cell.attrib[STYLE_TAG])

            cellid = ''.join(random.choice(string.ascii_lowercase) for i in range(100))

            for r in range(rowspan):
                for c in range(colspan):
                    dimens = {
                        TABLECELL_DIMENS_HEIGHT: height,
                        TABLECELL_DIMENS_WIDTH: width,
                        TABLECELL_DIMENS_CELLID: cellid,
                        TABLECELL_DIMENS_CELL: cell,
                    }

                    if dimenrowidx >= len(coldimens):
                        coldimens.append([])

                    while dimencolidx > len(coldimens[dimenrowidx]) - 1:
                        coldimens[dimenrowidx].append('')

                    coldimens[dimenrowidx][dimencolidx] = dimens

                    if r == rowspan - 1:
                        dimencolidx += 1
                dimenrowidx += 1

    return coldimens


def docx_build_table_rows_cols(docx, coldimens):
    numtablerows = len(coldimens)
    numtablecols = len(coldimens[0])
    docxtable = docx.add_table(rows=numtablerows, cols=numtablecols)
    default_height = Pt(numtablerows * 0.01)
    default_width = Pt(numtablecols * 0.01)

    cellid_map = dict()
    for rowidx in range(numtablerows):
        for colidx in range(numtablecols):
            dimens = coldimens[rowidx][colidx]
            cellid = dimens[TABLECELL_DIMENS_CELLID]

            if cellid not in cellid_map:
                docxtable.cell(rowidx, colidx).height = dimens[TABLECELL_DIMENS_HEIGHT] if dimens[
                                                                                               TABLECELL_DIMENS_HEIGHT] != -1 else default_height
                docxtable.cell(rowidx, colidx).width = dimens[TABLECELL_DIMENS_WIDTH] if dimens[
                                                                                             TABLECELL_DIMENS_WIDTH] != -1 else default_width
                cellid_map[cellid] = docxtable.cell(rowidx, colidx)
            else:
                cellid_map[cellid] = docxtable.cell(rowidx, colidx).merge(cellid_map[cellid])

    return docxtable


def docx_delete_paragraph(paragraph):
    p = paragraph._element

    if p.getparent() is not None:
        p.getparent().remove(p)
        p._p = p._element = None


def docx_cleanup_empty_parent_para(paragraph):
    if paragraph is not None and paragraph.text == '' and not docx_have_hyperlink(paragraph):
        docx_delete_paragraph(paragraph)


def docx_copy_paragraph(doc, paragraph, para_style=BODY_STYLE):
    old_paragraph = paragraph
    new_paragraph = doc.add_paragraph('', style=para_style)

    if old_paragraph is not None:
        new_paragraph.style = deepcopy(old_paragraph.style)
        new_paragraph.paragraph_format.alignment = old_paragraph.paragraph_format.alignment

    return new_paragraph


def docx_copy_run(paragraph, run, text='', run_style=''):
    old_run = run

    if old_run is not None:
        new_run = paragraph.add_run('')
        new_run.style = deepcopy(old_run.style)
        new_run.font.bold = old_run.font.bold
        new_run.font.cs_bold = old_run.font.cs_bold
        new_run.font.italic = old_run.font.italic
        new_run.font.cs_italic = old_run.font.cs_italic
        new_run.font.underline = old_run.font.underline
        new_run.font.superscript = old_run.font.superscript
        new_run.font.subscript = old_run.font.subscript
        new_run.font.color.rgb = old_run.font.color.rgb
        new_run.font.size = old_run.font.size
        new_run.text = new_run.text + text
    else:
        new_run = old_run
        if text is not '':
            paragraph.add_run(text, style=run_style)

    return new_run


def docx_resave_bad_image(image_filename):
    ImageFile.LOAD_TRUNCATED_IMAGES = True

    with open(image_filename, 'rb') as f:
        image_file = f
        image = Image.open(image_file)

        # next 3 lines strip exif
        data = list(image.getdata())
        image_without_exif = Image.new(image.mode, image.size)

    image_without_exif.putdata(data)
    image_without_exif.save(image_filename)


def docx_add_picture(image_filename, doc_or_run, image_caption='', is_paragraph=True, directory=''):
    try:
        doc_or_run.add_picture(image_filename, width=DEFAULT_IMAGE_WIDTH)
    except Exception:
        try:
            docx_resave_bad_image(image_filename)
            doc_or_run.add_picture(image_filename, width=DEFAULT_IMAGE_WIDTH)
        except Exception:
            write_error(directory, error='Image Error', exception=traceback.format_exc())
            doc_or_run.add_picture(GLOBALS[GLOBAL_LOGO_FILENAME], width=DEFAULT_IMAGE_WIDTH)

    if is_paragraph:
        doc_or_run.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

        if image_caption is not '':
            doc_or_run.paragraphs[-1].runs[0].add_break(WD_BREAK.LINE)
            doc_or_run.paragraphs[-1].add_run(DEFAULT_CAPTION_PADDING + image_caption, style=RUN_CAPTION_STYLE)


def docx_get_inline_image_prefix(idx):
    return 'INLINE_{}'.format(idx)
