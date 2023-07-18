from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx2pdf import convert
# from shared import docxtopdf
from shared.constants import *
from shared.docx_body_table import docx_build_body
from shared.docx_helpers import docx_add_bold, docx_add_underline, docx_apply_hyperlink_style, \
    docx_add_italic, docx_add_picture, docx_get_others_text
from shared.globals import GLOBALS
from shared.parse_helpers import parse_fetch_image


def docx_init_styles(styles):
    title_style = styles.add_style(TITLE_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = FONT_TNR
    title_style.font.size = Pt(14)
    docx_add_bold(title_style)

    datetime_style = styles.add_style(DATETIME_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    datetime_style.font.name = FONT_TNR
    datetime_style.font.size = Pt(10)

    caption_style = styles.add_style(CAPTION_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    caption_style.font.name = FONT_TNR
    caption_style.font.size = Pt(10)
    docx_add_italic(caption_style)

    run_caption_style = styles.add_style(RUN_CAPTION_STYLE, WD_STYLE_TYPE.CHARACTER)
    run_caption_style.font.name = FONT_TNR
    run_caption_style.font.size = Pt(10)
    docx_add_italic(run_caption_style)

    body_style = styles.add_style(BODY_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = FONT_TNR
    body_style.font.size = Pt(12)

    run_body_style = styles.add_style(RUN_BODY_STYLE, WD_STYLE_TYPE.CHARACTER)
    run_body_style.base_style = styles[BODY_STYLE]

    list_bullet_style = styles.add_style(LIST_BULLET_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    list_bullet_style.base_style = styles['List Bullet']
    list_bullet_style.font.name = FONT_TNR
    list_bullet_style.font.size = Pt(12)

    run_list_bullet_style = styles.add_style(RUN_LIST_BULLET_STYLE, WD_STYLE_TYPE.CHARACTER)
    run_list_bullet_style.base_style = styles[LIST_BULLET_STYLE]

    list_number_style = styles.add_style(LIST_NUMBER_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    list_number_style.base_style = styles['List Number']
    list_number_style.font.name = FONT_TNR
    list_number_style.font.size = Pt(12)

    run_list_number_style = styles.add_style(RUN_LIST_NUMBER_STYLE, WD_STYLE_TYPE.CHARACTER)
    run_list_number_style.base_style = styles[LIST_NUMBER_STYLE]

    more_resources_title_style = styles.add_style(MORE_RESOURCES_TITLE_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    more_resources_title_style.base_style = styles[BODY_STYLE]
    docx_add_bold(more_resources_title_style)

    more_resources_link_style = styles.add_style(MORE_RESOURCES_LINK_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    more_resources_link_style.font.name = FONT_TNR
    more_resources_link_style.font.size = Pt(10)
    docx_add_underline(more_resources_link_style)

    table_style = styles.add_style(TABLE_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    table_style.base_style = styles[BODY_STYLE]

    run_table_style = styles.add_style(RUN_TABLE_STYLE, WD_STYLE_TYPE.CHARACTER)
    run_table_style.base_style = styles[TABLE_STYLE]

    run_link_style = styles.add_style(RUN_LINK_STYLE, WD_STYLE_TYPE.CHARACTER)
    run_link_style.base_style = styles[RUN_BODY_STYLE]
    docx_apply_hyperlink_style(run_link_style)


"""
https://www.mindef.gov.sg/web/wcm/connect/mindef/mindef-content/home?siteAreaName=&srv=cmpnt&selectedCategories=news-releases&cmpntid=dcb39e68-0637-4383-b587-29be9bb9bea5&source=library&cache=none&contentcache=none&connectorcache=none&wcm_page.MENU-latest-releases=3
"""


def docx_build(save_filename, filename_prefix, directory,
               title, datetime_str, images, body, article_type,
               others_text=[], others_link=[], others_type=[]):
    print('Building DOCX: {}, {}'.format(save_filename, title))
    doc = Document()
    docx_init_styles(doc.styles)

    logo_at_top = GLOBALS[GLOBAL_SPEECH_LOGO_FILENAME] if SPEECHES_TYPE == article_type else GLOBALS[GLOBAL_LOGO_FILENAME]
    doc.add_picture(logo_at_top, width=DEFAULT_IMAGE_WIDTH)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph(title, style=TITLE_STYLE)

    doc.add_paragraph(datetime_str, style=DATETIME_STYLE)

    for idx, img in enumerate(images):
        img_filename = parse_fetch_image(img[PARSE_IMAGE_LINK], idx, filename_prefix, directory)
        images[idx][PARSE_IMAGE_LINK] = img_filename
    num_images = len(images)

    if num_images > 0:
        image_filename = images[0][PARSE_IMAGE_LINK]
        image_caption = images[0][PARSE_IMAGE_CAPTION]
        docx_add_picture(image_filename, doc, directory=directory, image_caption=image_caption)

    docx_build_body(body, doc, directory=directory, filename_prefix=filename_prefix)

    for i in range(1, num_images):
        image_filename = images[i][PARSE_IMAGE_LINK]
        image_caption = images[i][PARSE_IMAGE_CAPTION]
        docx_add_picture(image_filename, doc, directory=directory, image_caption=image_caption)

    num_texts = len(others_text)
    num_links = len(others_link)
    num_overall = num_texts if num_texts > num_links else num_links

    if num_overall > 0:
        type_map = dict()
        for i in range(num_overall):
            ol_type = others_type[i]
            type_map[ol_type] = type_map[ol_type] + 1 if ol_type in type_map else 1

        if len(type_map.keys()) > 1:
            others_title = MORE_RESOURCES_TITLE
        elif NEWS_RELEASES_TYPE in type_map:
            others_title = NEWS_RELEASE_TITLE if type_map[NEWS_RELEASES_TYPE] == 1 else NEWS_RELEASES_TITLE
        elif SPEECHES_TYPE in type_map:
            others_title = SPEECH_TITLE if type_map[SPEECHES_TYPE] == 1 else SPEECHES_TITLE
        elif OTHERS_TYPE in type_map:
            others_title = FACT_SHEET_TITLE if type_map[OTHERS_TYPE] == 1 else FACT_SHEETS_TITLE
        else:
            others_title = MORE_RESOURCES_TITLE

        doc.add_paragraph(others_title, style=MORE_RESOURCES_TITLE_STYLE)

        for i in range(num_overall):
            ol_text = others_text[i]
            ol_link = others_link[i]

            if ol_link != ERROR and ol_link != NOT_SUPPORTED:
                doc.add_paragraph(docx_get_others_text(ol_text, ol_link), style=LIST_BULLET_STYLE)

    save_path = os.path.join(directory, save_filename)
    print('Saving to DOCX: {}'.format(save_path))
    doc.save(save_path)

    with GLOBALS[GLOBAL_SAVE_PDF_COUNTER].get_lock():
        docx_input_path = save_path
        pdf_output_path = '{}{}'.format(save_path[:-4], 'pdf')
        print('Saving to PDF: {} --> {}'.format(docx_input_path, pdf_output_path))
        convert(input_path=docx_input_path, output_path=pdf_output_path)
        # docxtopdf.convert_to(folder=directory, source=save_path)

    return save_filename
