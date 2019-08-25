import re

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from lxml import etree
from shared.constants import RUN_BODY_STYLE, BODY_STYLE, TABLE_STYLE, RUN_LIST_BULLET_STYLE, RUN_LIST_NUMBER_STYLE, \
    HANDLED_TAGS, HREF_ATTRIB, LIST_BULLET_STYLE, LIST_NUMBER_STYLE, \
    LIST_TAGS, BOLD_TAGS, ITALIC_TAGS, PARAGRAPH_TAGS, REQ_NEW_PARA_TAGS, RUN_TABLE_STYLE, LIST_TYPE_UNORDERED, \
    LIST_TYPE_ORDERED, INLINE_TAGS, INLINE_IMAGE_IDX, SRC_ATTRIB, HEADING_TAGS, TABLECELL_DIMENS_CELL, \
    TABLECELL_DIMENS_CELLID, STYLE_TAG, LI_TAG, A_TAG, UL_TAG, OL_TAG, BR_TAG, HR_TAG, IMG_TAG, U_TAG, SUB_TAG, SUP_TAG, \
    P_TAG, NBSP, TABLE_TAG
from shared.docx_helpers import docx_add_bold, docx_add_hyperlink, docx_add_italic, docx_add_subscript, \
    docx_add_superscript, \
    docx_add_underline, docx_build_table_rows_cols, docx_delete_paragraph, \
    docx_add_heading, docx_get_coldimensions, docx_apply_text_align, docx_add_listitem, docx_cleanup_empty_parent_para, \
    docx_copy_paragraph, docx_copy_run, docx_add_picture, docx_get_inline_image_prefix
from shared.parse_helpers import parse_cleanup, parse_fetch_image, parse_append_hostname
from shared.writers import write_error, write_debug


def docx_build_table_cells(coldimens, docxtable, directory='', filename_prefix='', img_idx_obj=dict()):
    visit_map = dict()

    for rowidx in range(len(coldimens)):
        for colidx in range(len(coldimens[0])):
            dimens = coldimens[rowidx][colidx]
            cellid = dimens[TABLECELL_DIMENS_CELLID]

            if cellid in visit_map:
                continue
            visit_map[cellid] = True
            htmlcell = dimens[TABLECELL_DIMENS_CELL]
            docxcell = docxtable.cell(rowidx, colidx)
            docx_delete_paragraph(docxcell.paragraphs[0])
            write_debug(directory=directory, msg='TABLE - NEW PARA')
            para_to_use = docxcell.add_paragraph('', style=TABLE_STYLE)
            docx_apply_text_align(para_to_use, htmlcell.attrib)
            tablecol_text = htmlcell.text

            if tablecol_text is not None and tablecol_text != '':
                tablecol_text = parse_cleanup(htmlcell.text, dont_trim=True)
                para_to_use.text = tablecol_text
            if len(htmlcell.getchildren()) > 0:
                docx_build_body(htmlcell, doc=docxcell,
                                directory=directory, filename_prefix=filename_prefix,
                                paragraph=para_to_use,
                                parent_para_style=TABLE_STYLE, parent_run_style=RUN_TABLE_STYLE,
                                inline_img_idx_obj=img_idx_obj)


def docx_build_table(tablebody, doc, directory='', filename_prefix='', img_idx_obj=dict()):
    write_debug(directory=directory, msg='Build Table')
    coldimens = docx_get_coldimensions(tablebody)
    docxtable = docx_build_table_rows_cols(doc, coldimens)
    docxtable.alignment = WD_TABLE_ALIGNMENT.CENTER
    docx_build_table_cells(coldimens, docxtable,
                           directory=directory, filename_prefix=filename_prefix,
                           img_idx_obj=img_idx_obj)


def docx_build_li_para(doc, child_attrib, paragraph, para_style, directory='', debugmsg=''):
    write_debug(directory=directory, msg=debugmsg)

    if paragraph is None:
        li_para = doc.add_paragraph('', style=LIST_BULLET_STYLE)
        docx_apply_text_align(li_para, child_attrib)
    else:
        li_para = docx_copy_paragraph(doc, paragraph, para_style=para_style)
        docx_cleanup_empty_parent_para(paragraph)

    return li_para


def docx_build_li_run(paragraph, list_type, run_style):
    docx_add_listitem(paragraph, list_type)

    return paragraph.add_run('', style=run_style)


def docx_build_body(body, doc=None,
                    directory='', filename_prefix='',
                    paragraph=None, run=None,
                    parent_para_style=BODY_STYLE, parent_run_style=RUN_BODY_STYLE,
                    inline_img_idx_obj=dict()):
    children = body.getchildren()
    num_children = len(children)

    for idx in range(num_children):
        child_ele = children[idx]
        tag = child_ele.tag

        if tag is etree.Comment:
            write_error(directory, error='NOT HANDLED COMMENT')
            continue

        tag = tag.lower()

        if tag not in HANDLED_TAGS:
            write_error(directory, error='NOT HANDLED TAG: {}'.format(tag))
            continue
        if tag == STYLE_TAG:
            write_error(directory, error='STYLE TEXT: {}'.format(child_ele.text))
            continue
        elif tag == TABLE_TAG:
            docx_build_table(child_ele, doc, directory=directory, filename_prefix=filename_prefix,
                             img_idx_obj=inline_img_idx_obj)
            continue

        para_to_use = paragraph

        if para_to_use is None:
            write_debug(directory=directory, msg='NEW PARA')
            para_to_use = doc.add_paragraph('', style=parent_para_style)
            docx_apply_text_align(para_to_use, child_ele.attrib)

        raw_before_text = child_ele.text
        raw_after_text = child_ele.tail

        # Empty paragraph that acts as a separator
        if tag == P_TAG and raw_before_text == NBSP:
            child_ele.text = ' '
            raw_before_text = child_ele.text

        # Handle nbsp; before inline child
        if raw_before_text is not None and raw_before_text.endswith(NBSP) \
                and len(child_ele.getchildren()) > 0 \
                and child_ele.getchildren()[0].tag is not etree.Comment \
                and child_ele.getchildren()[0].tag in INLINE_TAGS:
            child_ele.text = re.sub('{}$'.format(NBSP), ' ', raw_before_text)
            raw_before_text = child_ele.text

        # Handle nbsp; after inline child
        if raw_after_text is not None and tag in INLINE_TAGS:
            if raw_after_text.startswith(NBSP):
                child_ele.tail = re.sub('^{}'.format(NBSP), ' ', child_ele.tail)
            if raw_after_text.endswith(NBSP):
                child_ele.tail = re.sub('{}$'.format(NBSP), ' ', child_ele.tail)
            raw_after_text = child_ele.tail

        before_text = parse_cleanup(raw_before_text if raw_before_text is not None else '', dont_trim=True)
        after_text = parse_cleanup(raw_after_text if raw_after_text is not None else '', dont_trim=True)
        before_run = run
        run_style = parent_run_style
        write_debug(directory=directory,
                    msg='INFO BUILD BODY {}, |{}|, |{}|, {}, para style: {}'.format(tag,
                                                                                    before_text, after_text,
                                                                                    para_to_use.paragraph_format.alignment,
                                                                                    para_to_use.style.name))

        ###
        # Adding of run and paragraph if necessary and various conditions for run
        ###
        if tag == A_TAG and HREF_ATTRIB in child_ele.attrib \
                and child_ele.attrib[HREF_ATTRIB] is not None \
                and child_ele.attrib[HREF_ATTRIB] != '':
            url = child_ele.attrib[HREF_ATTRIB]
            before_run = docx_add_hyperlink(para_to_use, url, before_text, old_run=before_run)
        elif tag == LI_TAG and run_style == RUN_LIST_BULLET_STYLE:
            para_to_use = docx_build_li_para(doc, child_ele.attrib,
                                             paragraph, para_style=LIST_BULLET_STYLE,
                                             directory=directory, debugmsg='LI BULLET TAG - NEW PARA & RUN')
            before_run = docx_build_li_run(para_to_use, LIST_TYPE_UNORDERED, run_style=run_style)
        elif tag == LI_TAG and run_style == RUN_LIST_NUMBER_STYLE:
            para_to_use = docx_build_li_para(doc, child_ele.attrib,
                                             paragraph, para_style=LIST_NUMBER_STYLE,
                                             directory=directory, debugmsg='LI NUMBER TAG - NEW PARA & RUN')
            before_run = docx_build_li_run(para_to_use, LIST_TYPE_ORDERED, run_style=run_style)
        elif tag in PARAGRAPH_TAGS:
            write_debug(directory=directory, msg='PARA TAG - NEW RUN')
            para_to_use = docx_copy_paragraph(doc, paragraph, para_style=parent_para_style)
            docx_cleanup_empty_parent_para(paragraph)
            paragraph = para_to_use
            docx_apply_text_align(para_to_use, child_ele.attrib)
            before_run = docx_copy_run(para_to_use, run)
            before_run = para_to_use.add_run('', style=run_style) if before_run is None else before_run
        elif before_run is None:
            write_debug(directory=directory, msg='NEW RUN')
            if tag == UL_TAG:
                run_style = RUN_LIST_BULLET_STYLE
            elif tag == OL_TAG:
                run_style = RUN_LIST_NUMBER_STYLE
            before_run = para_to_use.add_run('', style=run_style)
        else:
            before_run = docx_copy_run(para_to_use, run)
            before_run.text = ''
            write_debug(directory=directory,
                        msg='COPY RUN {}, {}, {}, {}'.format(before_run.font.bold, before_run.font.italic,
                                                             before_run.font.underline, before_run.font.color.rgb))

        # Insert Text, Break or Image
        if tag == BR_TAG:
            before_run.add_break(WD_BREAK.LINE)
        elif tag == HR_TAG:
            before_run.add_break(WD_BREAK.PAGE)
        elif tag == IMG_TAG:
            image_link = parse_append_hostname(child_ele.attrib[SRC_ATTRIB]) if SRC_ATTRIB in child_ele.attrib else ''
            print(image_link)
            inline_img_idx = inline_img_idx_obj[INLINE_IMAGE_IDX] if INLINE_IMAGE_IDX in inline_img_idx_obj else 0
            image_filename = parse_fetch_image(url=image_link, idx=docx_get_inline_image_prefix(inline_img_idx),
                                               directory=directory,
                                               filename_prefix=filename_prefix) if image_link is not '' else ''
            inline_img_idx_obj[INLINE_IMAGE_IDX] = inline_img_idx + 1
            docx_add_picture(image_filename, before_run, directory=directory, is_paragraph=False)
        elif tag != A_TAG and before_text != '':
            before_run.text = before_text

        # Style Text
        if tag in BOLD_TAGS:
            docx_add_bold(before_run)
        elif tag in ITALIC_TAGS:
            docx_add_italic(before_run)
        elif tag == U_TAG:
            docx_add_underline(before_run)
        elif tag == SUB_TAG:
            docx_add_subscript(before_run)
        elif tag == SUP_TAG:
            docx_add_superscript(before_run)
        elif tag in HEADING_TAGS:
            docx_add_heading(before_run)

        # Insert Text After Element
        if after_text is not '' and tag not in REQ_NEW_PARA_TAGS:
            if run is None:
                write_debug(directory=directory, msg='AFTER - NEW RUN')
                para_to_use.add_run(after_text, style=parent_run_style)
            else:
                write_debug(directory=directory, msg='AFTER - COPY RUN')
                docx_copy_run(para_to_use, run, text=after_text)

        # LIST TAG - Delete Paragraph if no text and only 1 run (this run)
        if tag in LIST_TAGS and before_run.text is '' and len(para_to_use.runs) == 1:
            docx_delete_paragraph(para_to_use)

        # Only pass run if its not a paragraph tag
        docx_build_body(body=child_ele, doc=doc,
                        directory=directory, filename_prefix=filename_prefix,
                        paragraph=para_to_use, run=before_run if tag != P_TAG else None,
                        parent_para_style=parent_para_style, parent_run_style=run_style,
                        inline_img_idx_obj=inline_img_idx_obj)

        # Add New Paragraph after closing list or new paragraph, assign new paragraph as the parent paragraph inheriting all the
        # properties. Other Conditions: Next Tag is not REQ_NEW_PARA_TAGS AND (Have After Text OR Not Last Child)
        next_idx = idx + 1
        is_not_last_child = next_idx < num_children
        next_tag = children[next_idx].tag.lower() if is_not_last_child and children[
            next_idx].tag is not etree.Comment else ''
        is_next_tag_req_new_para = is_not_last_child and next_tag not in REQ_NEW_PARA_TAGS

        if (after_text is not '' or is_next_tag_req_new_para) and tag in REQ_NEW_PARA_TAGS:
            write_debug(directory=directory, msg='AFTER CHILDREN - COPY PARA & RUN: {}, {}, {}'.format(tag, next_tag, after_text))
            paragraph = docx_copy_paragraph(doc, paragraph, para_style=parent_para_style)
            run = docx_copy_run(paragraph, run, text=after_text, run_style=parent_run_style)

    return []
