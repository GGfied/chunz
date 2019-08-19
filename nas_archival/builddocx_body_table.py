from copy import deepcopy

from constants import RUN_BODY_STYLE, BODY_STYLE, TABLE_STYLE, RUN_LIST_BULLET_STYLE, RUN_LIST_NUMBER_STYLE, \
    HANDLED_TAGS, HREF, LIST_BULLET_STYLE, LIST_NUMBER_STYLE, \
    LIST_TAGS, BOLD_TAGS, ITALIC_TAGS, PARAGRAPH_TAGS, REQ_NEW_PARA_TAGS, RUN_TABLE_STYLE, LIST_TYPE_UNORDERED, \
    LIST_TYPE_ORDERED
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx_helpers import docx_add_bold, docx_add_hyperlink, docx_add_italic, docx_add_subscript, docx_add_superscript, \
    docx_add_underline, is_heading_tag, docx_build_table_rows_cols, docx_delete_paragraph, \
    docx_add_heading, docx_get_coldimensions, docx_apply_text_align, docx_add_listitem, docx_cleanup_empty_parent_para
from parse_helpers import parse_cleanup


def docx_build_table_cells(coldimens, docxtable):
    visit_map = dict()

    for rowidx in range(len(coldimens)):
        for colidx in range(len(coldimens[0])):
            dimens = coldimens[rowidx][colidx]
            cellid = dimens['cellid']

            if cellid in visit_map:
                continue
            visit_map[cellid] = True
            tablecol = dimens['col']
            docxcell = docxtable.cell(rowidx, colidx)
            docx_delete_paragraph(docxcell.paragraphs[0])
            # print('TABLE - NEW PARA')
            para_to_use = docxcell.add_paragraph('', style=TABLE_STYLE)
            docx_apply_text_align(para_to_use, tablecol.attrib)
            tablecol_text = tablecol.text

            if tablecol_text is not None and tablecol_text != '':
                tablecol_text = parse_cleanup(tablecol.text, dont_trim=True)
                para_to_use.text = tablecol_text
            if len(tablecol.getchildren()) > 0:
                docx_build_body(tablecol, doc=docxcell, paragraph=para_to_use, parent_para_style=TABLE_STYLE,
                                parent_run_style=RUN_TABLE_STYLE)


def docx_build_table(tablebody, doc):
    # print('Build Table')
    coldimens = docx_get_coldimensions(tablebody)
    docxtable = docx_build_table_rows_cols(doc, coldimens)
    docxtable.alignment = WD_TABLE_ALIGNMENT.CENTER
    docx_build_table_cells(coldimens, docxtable)


def copy_paragraph(doc, paragraph, para_style=BODY_STYLE):
    old_paragraph = paragraph
    new_paragraph = doc.add_paragraph('', style=para_style)

    if old_paragraph is not None:
        new_paragraph.style = deepcopy(old_paragraph.style)
        new_paragraph.paragraph_format.alignment = old_paragraph.paragraph_format.alignment

    return new_paragraph


def copy_run(paragraph, run, text='', run_style=''):
    old_run = run

    if old_run is not None:
        new_run = paragraph.add_run('')
        new_run.style = deepcopy(old_run.style)
        new_run.font.bold = old_run.font.bold
        new_run.font.cs_bold = old_run.font.cs_bold
        new_run.font.italic = old_run.font.italic
        new_run.font.cs_italic = old_run.font.cs_italic
        new_run.font.underline = old_run.font.underline
        new_run.font.superscript = old_run.font.superscript
        new_run.font.subscript = old_run.font.subscript
        new_run.font.color.rgb = old_run.font.color.rgb
        new_run.font.size = old_run.font.size
        new_run.text = new_run.text + text
    else:
        new_run = old_run
        if text is not '':
            paragraph.add_run(text, style=run_style)

    return new_run


def docx_build_body(body, doc=None, paragraph=None, run=None, parent_para_style=BODY_STYLE,
                    parent_run_style=RUN_BODY_STYLE):
    children = body.getchildren()
    num_children = len(children)

    for idx in range(num_children):
        e = children[idx]
        tag = e.tag.lower()

        if tag not in HANDLED_TAGS:
            print('NOT HANDLED TAG', tag)

        if tag == 'table':
            docx_build_table(e, doc)
            continue

        para_to_use = paragraph
        if para_to_use is None:
            # print('NEW PARA', e, e.attrib)
            para_to_use = doc.add_paragraph('', style=parent_para_style)
            docx_apply_text_align(para_to_use, e.attrib)

        if tag == 'p' and e.text == '\xa0':
            e.text = ' '

        before_text = parse_cleanup(e.text if e.text is not None else '', dont_trim=True)
        after_text = parse_cleanup(e.tail if e.tail is not None else '', dont_trim=True)
        before_run = run
        run_style = parent_run_style
        # print('INFO', tag, '|' + before_text, '|' + after_text,
        #       para_to_use.paragraph_format.alignment)

        # Start of any element
        if tag == 'a' and HREF in e.attrib and e.attrib[HREF] is not None and e.attrib[HREF] != '':
            url = e.attrib[HREF]
            before_run = docx_add_hyperlink(para_to_use, url, before_text, old_run=before_run)
        elif tag == 'li' and run_style == RUN_LIST_BULLET_STYLE:
            # print('LI BULLET TAG - NEW PARA & RUN')
            if paragraph is None:
                para_to_use = doc.add_paragraph('', style=LIST_BULLET_STYLE)
                docx_apply_text_align(para_to_use, e.attrib)
            else:
                para_to_use = copy_paragraph(doc, paragraph, para_style=LIST_BULLET_STYLE)
                docx_cleanup_empty_parent_para(paragraph)
            docx_add_listitem(para_to_use, LIST_TYPE_UNORDERED)
            before_run = para_to_use.add_run('', style=run_style)
        elif tag == 'li' and run_style == RUN_LIST_NUMBER_STYLE:
            # print('LI NUMBER TAG - NEW PARA & RUN')
            if paragraph is None:
                para_to_use = doc.add_paragraph('', style=LIST_NUMBER_STYLE)
                docx_apply_text_align(para_to_use, e.attrib)
            else:
                para_to_use = copy_paragraph(doc, paragraph, para_style=LIST_NUMBER_STYLE)
                docx_cleanup_empty_parent_para(paragraph)
            docx_add_listitem(para_to_use, LIST_TYPE_ORDERED)
            before_run = para_to_use.add_run('', style=run_style)
        elif tag in PARAGRAPH_TAGS:
            # print('PARA TAG - NEW RUN')
            para_to_use = copy_paragraph(doc, paragraph, para_style=parent_para_style)
            docx_cleanup_empty_parent_para(paragraph)
            paragraph = para_to_use
            docx_apply_text_align(para_to_use, e.attrib)
            before_run = copy_run(para_to_use, run)
            before_run = para_to_use.add_run('', style=run_style) if before_run is None else before_run
        elif before_run is None:
            # print('NEW RUN')
            if tag == 'ul':
                run_style = RUN_LIST_BULLET_STYLE
            elif tag == 'ol':
                run_style = RUN_LIST_NUMBER_STYLE
            before_run = para_to_use.add_run('', style=run_style)
        else:
            before_run = copy_run(para_to_use, run)
            before_run.text = ''
            # print('COPY RUN',
            #       before_run.font.bold,
            #       before_run.font.italic)

        # Insert Text or Break
        if tag == 'br':
            before_run.add_break(WD_BREAK.LINE)
        elif tag == 'hr':
            before_run.add_break(WD_BREAK.PAGE)
        elif tag != 'a' and before_text != '':
            before_run.text = before_text

        # Style Text
        if tag in BOLD_TAGS:
            docx_add_bold(before_run)
        elif tag in ITALIC_TAGS:
            docx_add_italic(before_run)
        elif tag == 'u':
            docx_add_underline(before_run)
        elif tag == 'sub':
            docx_add_subscript(before_run)
        elif tag == 'sup':
            docx_add_superscript(before_run)
        elif is_heading_tag(tag):
            docx_add_heading(before_run)

        # Insert Text After Element
        if after_text is not '' and tag not in REQ_NEW_PARA_TAGS:
            if run is None:
                # print('AFTER - NEW RUN')
                para_to_use.add_run(after_text, style=parent_run_style)
            else:
                # print('AFTER - COPY RUN')
                copy_run(para_to_use, run, text=after_text)

        # LIST TAG - Delete Paragraph if no text and only 1 run (this run)
        if tag in LIST_TAGS and before_run.text is '' and len(para_to_use.runs) == 1:
            docx_delete_paragraph(para_to_use)

        # Only pass run if its not a paragraph tag
        docx_build_body(e, doc=doc, paragraph=para_to_use,
                        run=before_run if tag != 'p' else None,
                        parent_para_style=parent_para_style,
                        parent_run_style=run_style)

        # Add New Paragraph after closing list or new paragraph, assign new paragraph as the parent paragraph inheriting all the
        # properties. Other Conditions: Next Tag is not REQ_NEW_PARA_TAGS AND (Have After Text OR Not Last Child)
        next_idx = idx + 1
        is_not_last_child = next_idx < num_children
        next_tag = children[next_idx].tag.lower() if is_not_last_child else ''
        is_next_tag_req_new_para = is_not_last_child and next_tag not in REQ_NEW_PARA_TAGS

        if (after_text is not '' or is_next_tag_req_new_para) and tag in REQ_NEW_PARA_TAGS:
            # print('AFTER CHILDREN - COPY PARA & RUN')
            paragraph = copy_paragraph(doc, paragraph, para_style=parent_para_style)
            run = copy_run(paragraph, run, text=after_text, run_style=parent_run_style)

    return []
