import random
import re
import string

import docx
from constants import WIDTH, COLSPAN, RUN_LINK_STYLE, STYLE, ROWSPAN, HEIGHT, TEXT_ALIGN_IN_STYLE_RE, \
    HEIGHT_IN_STYLE_RE, WIDTH_IN_STYLE_RE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def is_heading_tag(tag):
    return tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']


def docx_add_bold(style):
    style.font.bold = True
    style.font.cs_bold = True


def docx_add_italic(style):
    style.font.italic = True
    style.font.cs_italic = True


def docx_add_underline(style):
    style.font.underline = True


def docx_add_superscript(style):
    style.font.superscript = True


def docx_add_subscript(style):
    style.font.subscript = True


def docx_apply_hyperlink_style(style):
    style.font.color.rgb = RGBColor(0, 0, 255)
    style.font.underline = True


def docx_add_heading(style):
    style.font.size = Pt(12)
    docx_add_bold(style)


def docx_apply_center_align(style):
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def docx_apply_text_align(para, htmlattrib):
    whichalign = 'left'

    if STYLE in htmlattrib:
        print(htmlattrib[STYLE])
        res = re.search(TEXT_ALIGN_IN_STYLE_RE, htmlattrib[STYLE])
        whichalign = res.group(1).lower() if res else whichalign

    print(htmlattrib.keys(), whichalign)

    if whichalign == 'right':
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif whichalign == 'left':
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif whichalign == 'center':
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


"""
Source: https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx
Source 2: https://github.com/python-openxml/python-docx/issues/384
Source 3: https://github.com/python-openxml/python-docx/issues/74
:param paragraph: The paragraph we are adding the hyperlink to.
:param url: A string containing the required url
:param text: The text displayed for the url
:return: The hyperlink object
"""


def docx_add_hyperlink(paragraph, url, text, old_run=None):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    new_run = paragraph.add_run('', style=RUN_LINK_STYLE) if old_run is None else old_run
    new_run.text = new_run.text + text
    docx_apply_hyperlink_style(new_run.style)

    hyperlink.append(new_run._r)

    paragraph._p.append(hyperlink)

    return new_run


"""
Source: https://github.com/python-openxml/python-docx/issues/217
"""


def docx_add_listitem(paragraph, list_type):
    p = paragraph._p  # access to xml paragraph element
    pPr = p.get_or_add_pPr()  # access paragraph properties
    numPr = docx.oxml.shared.OxmlElement('w:numPr')  # create number properties element
    numId = docx.oxml.shared.OxmlElement('w:numId')  # create numId element - sets bullet type
    numId.set(docx.oxml.ns.qn('w:val'), list_type)  # set list type/indentation
    numPr.append(numId)  # add bullet type to number properties list
    pPr.append(numPr)  # add number properties to paragraph


def get_height_in_style(style):
    res = re.search(HEIGHT_IN_STYLE_RE, style)

    return res.group(1) if res else -1


def get_width_in_style(style):
    res = re.search(WIDTH_IN_STYLE_RE, style)

    return res.group(1) if res else -1


def docx_get_coldimensions(tablebody):
    tablerows = tablebody.getchildren()
    is_first_element_tbody = len(tablerows) > 0 and tablerows[0].tag.lower() == 'tbody'
    tablerows = tablerows[0].getchildren() if is_first_element_tbody else tablerows
    coldimens = []

    for rowidx, row in enumerate(tablerows):
        if row.tag.lower() != 'tr':
            continue

        if rowidx >= len(coldimens):
            coldimens.append([])
        dimencolidx = 0

        for colidx, col in enumerate(row.getchildren()):
            dimenrowidx = rowidx
            if col.tag.lower() != 'td':
                continue

            while dimencolidx in coldimens[dimenrowidx]:
                print('Add', dimencolidx)
                dimencolidx += 1

            colspan = 1
            if COLSPAN in col.attrib:
                try:
                    colspan = int(col.attrib[COLSPAN])
                except ValueError:
                    colspan = 1

            width = -1
            if WIDTH in col.attrib:
                try:
                    width = int(col.attrib[WIDTH])
                except ValueError:
                    width = -1

            if width == -1 and STYLE in col.attrib:
                width = get_width_in_style(col.attrib[STYLE])

            rowspan = 1
            if ROWSPAN in col.attrib:
                try:
                    rowspan = int(col.attrib[ROWSPAN])
                except ValueError:
                    rowspan = 1

            height = -1
            if HEIGHT in col.attrib:
                try:
                    height = int(col.attrib[HEIGHT])
                except ValueError:
                    height = -1

            if height == -1 and STYLE in col.attrib:
                height = get_height_in_style(col.attrib[STYLE])

            cellid = ''.join(random.choice(string.ascii_lowercase) for i in range(1000))
            for r in range(rowspan):
                for c in range(colspan):
                    dimens = {
                        'height': height,
                        'width': width,
                        'cellid': cellid,
                        'col': col,
                    }

                    if dimenrowidx >= len(coldimens):
                        coldimens.append([])
                    if dimencolidx not in coldimens[dimenrowidx]:
                        coldimens[dimenrowidx].append(dimens)
                    else:
                        coldimens[dimenrowidx][dimencolidx] = dimens

                    if r == rowspan - 1:
                        dimencolidx += 1
                dimenrowidx += 1

    return coldimens


def docx_build_table_rows_cols(docx, coldimens):
    numtablerows = len(coldimens)
    numtablecols = len(coldimens[0])
    docxtable = docx.add_table(rows=numtablerows, cols=numtablecols)
    default_height = Pt(numtablerows * 0.01)
    default_width = Pt(numtablecols * 0.01)

    cellid_map = dict()
    for rowidx in range(numtablerows):
        for colidx in range(numtablecols):
            dimens = coldimens[rowidx][colidx]
            cellid = dimens['cellid']

            if cellid not in cellid_map:
                docxtable.cell(rowidx, colidx).height = 1#dimens['height']  if dimens['height'] != -1 else default_height
                docxtable.cell(rowidx, colidx).width = 1#dimens['width'] if dimens['width'] != -1 else default_width
                cellid_map[cellid] = docxtable.cell(rowidx, colidx)
            else:
                docxtable.cell(rowidx, colidx).merge(cellid_map[cellid])

    return docxtable


def docx_delete_paragraph(paragraph):
    p = paragraph._element

    if p.getparent() is not None:
        p.getparent().remove(p)
        p._p = p._element = None
